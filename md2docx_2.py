#!/usr/bin/env python3
"""
Enhanced Markdown to Beautiful DOCX Converter with Code Block Styling,
Inline Formatting, Table Support, and a Dedicated Table of Contents Page.

This script converts a Markdown file into a beautifully formatted DOCX file
with custom styles, page margins, header/footer, colors, typeface, and specialized
styling for code blocks, tables, and inline formatting. It also detects a "Table of Contents"
section (if present) and creates a dedicated index page with clickable-style anchor links.

Supported features:
  - Title and cover page
  - Dedicated index page (if a Table of Contents exists)
  - Headings (levels 1-6)
  - Fenced code blocks (using triple backticks) with shading and indentation
  - Tables with proper alignment and styling
  - Block quotes
  - Unordered lists (including basic nested lists)
  - Bold, italic, and bold-italic inline formatting
  - Regular paragraphs

Usage:
    python md_to_docx_beautiful.py input.md output.docx [--title "Document Title"]

Dependencies:
    - python-docx (install via pip install python-docx)
"""
import argparse
import os
import sys
import re
import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

INDENT_SPACES_PER_LEVEL = 2  # Number of spaces for one level of list indentation
LIST_INDENT_INCREMENT = Inches(0.25)  # How much to indent per list level
BASE_LIST_INDENT = Inches(0.25)      # Initial indent for any list item
HANGING_LIST_INDENT = Inches(-0.25)   # To make bullets hang

def set_document_margins(document):
    """
    Set custom page margins for the document.
    """
    for section in document.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def set_cell_border(cell, **kwargs):
    """
    Set cell border with more control over styling
    Usage:
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#4F81BD"},
            bottom={"sz": 12, "color": "#4F81BD"},
            left={"sz": 24, "val": "dashed", "shadow": "true"},
            right={"sz": 12, "val": "dashed"},
        )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            tag = 'w:{}'.format(edge)
            
            # check for tag existence, if none found, create one
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
                
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in kwargs[edge]:
                    element.set(qn('w:{}'.format(key)), str(kwargs[edge][key]))


def set_cell_shading(cell, color):
    """
    Apply shading/background color to a table cell.
    """
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_vertical_alignment(cell, alignment):
    """
    Set vertical alignment of content in a cell.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    valign = OxmlElement('w:vAlign')
    valign.set(qn('w:val'), alignment)
    tcPr.append(valign)


def set_custom_styles(document):
    primary_color_rgb = RGBColor(0x00, 0x5A, 0x9C)
    primary_color_hex = "005A9C"
    gray_color_hex = "A9A9A9"

    title_style = document.styles['Title']
    title_font = title_style.font
    title_font.name = 'Cambria'
    title_font.size = Pt(36)
    title_font.bold = True
    title_font.color.rgb = primary_color_rgb

    normal_style = document.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    heading_sizes = [24, 18, 14, 12, 11, 11]
    heading_space_after = [Pt(18), Pt(12), Pt(8), Pt(6), Pt(6), Pt(6)]
    for i in range(1, 7):
        style = document.styles[f'Heading {i}']
        font = style.font
        font.name = 'Calibri Light'
        font.size = Pt(heading_sizes[i-1])
        font.bold = False 
        font.color.rgb = primary_color_rgb
        
        p_fmt = style.paragraph_format
        p_fmt.space_after = heading_space_after[i-1]
        if i == 1:
            font.bold = True
            p_fmt.space_before = Pt(24)
            p_fmt.keep_with_next = True
            pPr = style.paragraph_format._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom_b = OxmlElement('w:bottom')
            bottom_b.set(qn('w:val'), 'single')
            bottom_b.set(qn('w:sz'), '4') 
            bottom_b.set(qn('w:space'), '1')
            bottom_b.set(qn('w:color'), primary_color_hex)
            pBdr.append(bottom_b)
            pPr.append(pBdr)
        elif i == 2:
            p_fmt.keep_with_next = True
            pPr = style.paragraph_format._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom_b = OxmlElement('w:bottom')
            bottom_b.set(qn('w:val'), 'single')
            bottom_b.set(qn('w:sz'), '2') 
            bottom_b.set(qn('w:space'), '1')
            bottom_b.set(qn('w:color'), primary_color_hex)
            pBdr.append(bottom_b)
            pPr.append(pBdr)
        elif i == 3:
            p_fmt.keep_with_next = True

    code_style_name = "Code"
    try:
        _ = document.styles[code_style_name]
    except KeyError:
        code_style = document.styles.add_style(code_style_name, 1)
        code_font = code_style.font
        code_font.name = "Consolas"
        code_font.size = Pt(9.5)
        code_font.color.rgb = RGBColor(0x20, 0x20, 0x20)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        code_style.paragraph_format.line_spacing = 1.0
    
    blockquote_style_name = "BlockQuote"
    try:
        _ = document.styles[blockquote_style_name]
    except KeyError:
        blockquote_style = document.styles.add_style(blockquote_style_name, 1)
        blockquote_font = blockquote_style.font
        blockquote_font.name = "Calibri"
        blockquote_font.size = Pt(10.5)
        blockquote_font.italic = True
        blockquote_font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        p_fmt = blockquote_style.paragraph_format
        p_fmt.left_indent = Inches(0.75)
        p_fmt.space_before = Pt(6)
        p_fmt.space_after = Pt(6)
        p_fmt.line_spacing = 1.15
        
        pPr = blockquote_style.paragraph_format._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left_border = OxmlElement('w:left')
        left_border.set(qn('w:val'), 'single')
        left_border.set(qn('w:sz'), '6') 
        left_border.set(qn('w:space'), '8') 
        left_border.set(qn('w:color'), primary_color_hex)
        pBdr.append(left_border)
        pPr.append(pBdr)
    
    # Add style for math equations
    math_style_name = "MathEquation"
    try:
        _ = document.styles[math_style_name]
    except KeyError:
        math_style = document.styles.add_style(math_style_name, 1)
        math_font = math_style.font
        math_font.name = "Cambria Math"
        math_font.size = Pt(11)
        math_font.italic = True
        math_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        p_fmt = math_style.paragraph_format
        p_fmt.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p_fmt.space_before = Pt(6)
        p_fmt.space_after = Pt(6)
        
    return code_style_name, blockquote_style_name, math_style_name, primary_color_rgb


def add_decorative_horizontal_line(doc, color_hex, size_pt, space_before_pt=6, space_after_pt=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after = Pt(space_after_pt)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), str(int(size_pt * 8))) # Size is in 1/8ths of a point
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), color_hex)
    pBdr.append(bottom_border)
    pPr.append(pBdr)


def add_header_footer(document, title="Document", author=""):
    primary_color_hex = "005A9C"
    gray_color_rgb = RGBColor(0xA9, 0xA9, 0xA9)
    gray_color_hex = "A9A9A9"

    for section in document.sections:
        header = section.header
        if header.is_linked_to_previous is False: 
            header.paragraphs[0].clear() 
            header_paragraph = header.paragraphs[0]
        elif not header.paragraphs:
            header_paragraph = header.add_paragraph()
        else:
            header_paragraph = header.paragraphs[0]
            header_paragraph.clear()

        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        pPr_header = header_paragraph._p.get_or_add_pPr()
        pBdr_header = OxmlElement('w:pBdr')
        bottom_border_header = OxmlElement('w:bottom')
        bottom_border_header.set(qn('w:val'), 'single')
        bottom_border_header.set(qn('w:sz'), '4')
        bottom_border_header.set(qn('w:space'), '1')
        bottom_border_header.set(qn('w:color'), primary_color_hex)
        pBdr_header.append(bottom_border_header)
        pPr_header.append(pBdr_header)
        
        header_run = header_paragraph.add_run(title)
        header_run.font.name = "Calibri"
        header_run.font.size = Pt(9)
        header_run.font.color.rgb = gray_color_rgb

        footer = section.footer
        if footer.is_linked_to_previous is False:
            footer.paragraphs[0].clear()
            footer_paragraph = footer.paragraphs[0]
        elif not footer.paragraphs:
            footer_paragraph = footer.add_paragraph()
        else:
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.clear()
            
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        pPr_footer = footer_paragraph._p.get_or_add_pPr()
        pBdr_footer = OxmlElement('w:pBdr')
        top_border_footer = OxmlElement('w:top')
        top_border_footer.set(qn('w:val'), 'single')
        top_border_footer.set(qn('w:sz'), '4') 
        top_border_footer.set(qn('w:color'), gray_color_hex)
        pBdr_footer.append(top_border_footer)
        pPr_footer.append(pBdr_footer)
        
        run = footer_paragraph.add_run()
        
        fldChar_page_begin = OxmlElement('w:fldChar')
        fldChar_page_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar_page_begin)
        
        instrText_page = OxmlElement('w:instrText')
        instrText_page.set(qn('xml:space'), 'preserve')
        instrText_page.text = "PAGE"
        run._r.append(instrText_page)
        
        fldChar_page_sep = OxmlElement('w:fldChar')
        fldChar_page_sep.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar_page_sep)

        fldChar_page_end = OxmlElement('w:fldChar')
        fldChar_page_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar_page_end)

        run.add_text(' of ')

        fldChar_numpages_begin = OxmlElement('w:fldChar')
        fldChar_numpages_begin.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar_numpages_begin)

        instrText_numpages = OxmlElement('w:instrText')
        instrText_numpages.set(qn('xml:space'), 'preserve')
        instrText_numpages.text = "NUMPAGES"
        run._r.append(instrText_numpages)
        
        fldChar_numpages_sep = OxmlElement('w:fldChar')
        fldChar_numpages_sep.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar_numpages_sep)

        fldChar_numpages_end = OxmlElement('w:fldChar')
        fldChar_numpages_end.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar_numpages_end)
        
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.color.rgb = gray_color_rgb
        
        if author:
            author_run = footer_paragraph.add_run(f"  •  {author}")
            author_run.font.name = 'Calibri'
            author_run.font.size = Pt(8)
            author_run.font.italic = True
            author_run.font.color.rgb = gray_color_rgb


def apply_code_block_formatting(paragraph):
    """
    Apply formatting to a code block paragraph.
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), "F5F5F5")
    pPr.append(shd)
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.right_indent = Inches(0.2)
    
    pBdr = OxmlElement('w:pBdr')
    top_border = OxmlElement('w:top')
    top_border.set(qn('w:val'), 'single')
    top_border.set(qn('w:sz'), '2')
    top_border.set(qn('w:color'), 'D3D3D3')
    pBdr.append(top_border)
    
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '2')
    bottom_border.set(qn('w:color'), 'D3D3D3')
    pBdr.append(bottom_border)
    pPr.append(pBdr)


def process_math_equation(text):
    """
    Extract math equation from delimiters and return cleaned equation text.
    """
    # Remove delimiters ($ or $$) and trim whitespace
    if text.startswith('$$') and text.endswith('$$'):
        return text[2:-2].strip()
    elif text.startswith('$') and text.endswith('$'):
        return text[1:-1].strip()
    return text


def latex_to_omml(latex_expr):
    """Convert LaTeX expressions to OMML (Office Math Markup Language)"""
    # Remove any $ delimiters if still present
    latex_expr = latex_expr.replace('$', '').strip()
    
    # This is a basic LaTeX to OMML converter
    omml_parts = []
    
    # Handle common LaTeX patterns
    import re
    
    def convert_fraction(match):
        numerator = match.group(1)
        denominator = match.group(2)
        return f'''<m:f>
            <m:fPr></m:fPr>
            <m:num><m:r><m:t>{numerator}</m:t></m:r></m:num>
            <m:den><m:r><m:t>{denominator}</m:t></m:r></m:den>
        </m:f>'''
    
    def convert_superscript(match):
        base = match.group(1)
        exp = match.group(2)
        return f'''<m:sSup>
            <m:sSupPr></m:sSupPr>
            <m:e><m:r><m:t>{base}</m:t></m:r></m:e>
            <m:sup><m:r><m:t>{exp}</m:t></m:r></m:sup>
        </m:sSup>'''
    
    def convert_subscript(match):
        base = match.group(1)
        sub = match.group(2)
        return f'''<m:sSub>
            <m:sSubPr></m:sSubPr>
            <m:e><m:r><m:t>{base}</m:t></m:r></m:e>
            <m:sub><m:r><m:t>{sub}</m:t></m:r></m:sub>
        </m:sSub>'''
    
    def convert_sqrt(match):
        content = match.group(1)
        return f'''<m:rad>
            <m:radPr></m:radPr>
            <m:deg></m:deg>
            <m:e><m:r><m:t>{content}</m:t></m:r></m:e>
        </m:rad>'''
    
    # Apply conversions in order
    latex_expr = re.sub(r'\\frac\{([^{}]+)\}\{([^{}]+)\}', convert_fraction, latex_expr)
    latex_expr = re.sub(r'([a-zA-Z0-9()]+)\^\{([^{}]+)\}', convert_superscript, latex_expr)
    latex_expr = re.sub(r'([a-zA-Z0-9()]+)\^([a-zA-Z0-9])', convert_superscript, latex_expr)
    latex_expr = re.sub(r'([a-zA-Z0-9]+)_\{([^{}]+)\}', convert_subscript, latex_expr)
    latex_expr = re.sub(r'([a-zA-Z0-9]+)_([a-zA-Z0-9])', convert_subscript, latex_expr)
    latex_expr = re.sub(r'\\sqrt\{([^{}]+)\}', convert_sqrt, latex_expr)
    
    # Convert Greek letters and symbols
    # Comprehensive LaTeX symbol to Unicode replacement dictionary
    symbol_replacements = {
        # Lowercase Greek letters
        '\\alpha': 'α', '\\beta': 'β', '\\gamma': 'γ', '\\delta': 'δ',
        '\\epsilon': 'ε', '\\varepsilon': 'ϵ', '\\zeta': 'ζ', '\\eta': 'η',
        '\\theta': 'θ', '\\vartheta': 'ϑ', '\\iota': 'ι', '\\kappa': 'κ',
        '\\lambda': 'λ', '\\mu': 'μ', '\\nu': 'ν', '\\xi': 'ξ',
        '\\omicron': 'ο', '\\pi': 'π', '\\varpi': 'ϖ', '\\rho': 'ρ',
        '\\varrho': 'ϱ', '\\sigma': 'σ', '\\varsigma': 'ς', '\\tau': 'τ',
        '\\upsilon': 'υ', '\\phi': 'φ', '\\varphi': 'ϕ', '\\chi': 'χ',
        '\\psi': 'ψ', '\\omega': 'ω',

        # Uppercase Greek letters
        '\\Gamma': 'Γ', '\\Delta': 'Δ', '\\Theta': 'Θ', '\\Lambda': 'Λ',
        '\\Xi': 'Ξ', '\\Pi': 'Π', '\\Sigma': 'Σ', '\\Upsilon': 'Υ',
        '\\Phi': 'Φ', '\\Psi': 'Ψ', '\\Omega': 'Ω',

        # Binary operators
        '\\pm': '±', '\\mp': '∓', '\\times': '×', '\\div': '÷',
        '\\ast': '∗', '\\star': '⋆', '\\circ': '∘', '\\bullet': '∙',
        '\\cdot': '·', '\\cap': '∩', '\\cup': '∪', '\\uplus': '⊎',
        '\\sqcap': '⊓', '\\sqcup': '⊔', '\\vee': '∨', '\\wedge': '∧',
        '\\setminus': '∖', '\\wr': '≀', '\\diamond': '⋄', '\\bigtriangleup': '△',
        '\\bigtriangledown': '▽', '\\triangleleft': '◁', '\\triangleright': '▷',
        '\\oplus': '⊕', '\\ominus': '⊖', '\\otimes': '⊗', '\\oslash': '⊘',
        '\\odot': '⊙', '\\bigcirc': '◯', '\\dagger': '†', '\\ddagger': '‡',
        '\\amalg': '⨿',

        # Relations
        '\\leq': '≤', '\\le': '≤', '\\geq': '≥', '\\ge': '≥', '\\neq': '≠',
        '\\ne': '≠', '\\equiv': '≡', '\\approx': '≈', '\\sim': '∼',
        '\\simeq': '≃', '\\cong': '≅', '\\propto': '∝', '\\in': '∈',
        '\\ni': '∋', '\\notin': '∉', '\\subset': '⊂', '\\supset': '⊃',
        '\\subseteq': '⊆', '\\supseteq': '⊇', '\\nsubseteq': '⊈', '\\nsupseteq': '⊉',
        '\\subsetneq': '⊊', '\\supsetneq': '⊋', '\\perp': '⊥', '\\parallel': '∥',
        '\\mid': '∣', '\\nmid': '∤', '\\vdash': '⊢', '\\dashv': '⊣',
        '\\models': '⊨', '\\smile': '⌣', '\\frown': '⌢', '\\asymp': '≍',
        '\\doteq': '≐', '\\bowtie': '⋈', '\\prec': '≺', '\\succ': '≻',
        '\\preceq': '≼', '\\succeq': '≽', '\\ll': '≪', '\\gg': '≫',
        '\\between': '≬', '\\not': '¬', '\\implies': '⇒', '\\impliedby': '⇐',

        # Arrows
        '\\leftarrow': '←', '\\gets': '←', '\\rightarrow': '→', '\\to': '→',
        '\\leftrightarrow': '↔', '\\uparrow': '↑', '\\downarrow': '↓',
        '\\updownarrow': '↕', '\\Leftarrow': '⇐', '\\Rightarrow': '⇒',
        '\\Leftrightarrow': '⇔', '\\Uparrow': '⇑', '\\Downarrow': '⇓',
        '\\Updownarrow': '⇕', '\\mapsto': '↦', '\\hookleftarrow': '↩',
        '\\hookrightarrow': '↪', '\\longleftarrow': '⟵', '\\longrightarrow': '⟶',
        '\\Longleftarrow': '⟸', '\\Longrightarrow': '⟹', '\\longleftrightarrow': '⟷',
        '\\Longleftrightarrow': '⟺',

        # Miscellaneous symbols
        '\\infty': '∞', '\\partial': '∂', '\\nabla': '∇', '\\aleph': 'ℵ',
        '\\imath': 'ı', '\\jmath': 'ȷ', '\\ell': 'ℓ', '\\wp': '℘',
        '\\Re': 'ℜ', '\\Im': 'ℑ', '\\top': '⊤', '\\bot': '⊥',
        '\\forall': '∀', '\\exists': '∃', '\\neg': '¬', '\\emptyset': '∅',
        '\\varnothing': '∅', '\\angle': '∠', '\\measuredangle': '∡',
        '\\triangle': '△', '\\Box': '□', '\\square': '□', '\\blacksquare': '■',
        '\\Diamond': '◇', '\\diamondsuit': '♦', '\\heartsuit': '♥',
        '\\clubsuit': '♣', '\\spadesuit': '♠', '\\flat': '♭', '\\natural': '♮',
        '\\sharp': '♯', '\\prime': '′', '\\doubleprime': '″', '\\backslash': '\\',
        '\\slash': '/', '\\degree': '°', '\\copyright': '©', '\\pounds': '£',
        '\\yen': '¥', '\\euro': '€', '\\S': '§', '\\P': '¶', '\\dag': '†',
        '\\ddag': '‡', '\\ldots': '…', '\\cdots': '⋯', '\\vdots': '⋮',
        '\\ddots': '⋱', '\\dots': '…', '\\aleph': 'ℵ', '\\hbar': 'ℏ',
        '\\ell': 'ℓ', '\\wp': '℘', '\\Re': 'ℜ', '\\Im': 'ℑ', '\\mho': '℧',
        '\\Finv': 'Ⅎ', '\\Game': '⅁', '\\eth': 'ð', '\\gimel': 'ℷ',
        '\\beth': 'ℶ', '\\daleth': 'ℸ', '\\Bbbk': '𝕜',

        # Set theory
        '\\cap': '∩', '\\cup': '∪', '\\setminus': '∖', '\\complement': '∁',
        '\\varnothing': '∅', '\\emptyset': '∅', '\\in': '∈', '\\notin': '∉',
        '\\ni': '∋', '\\subset': '⊂', '\\supset': '⊃', '\\subseteq': '⊆',
        '\\supseteq': '⊇', '\\nsubseteq': '⊈', '\\nsupseteq': '⊉',

        # Logic
        '\\land': '∧', '\\wedge': '∧', '\\lor': '∨', '\\vee': '∨',
        '\\neg': '¬', '\\lnot': '¬', '\\implies': '⇒', '\\impliedby': '⇐',
        '\\iff': '⇔', '\\therefore': '∴', '\\because': '∵',

        # Calculus
        '\\int': '∫', '\\iint': '∬', '\\iiint': '∭', '\\oint': '∮',
        '\\sum': '∑', '\\prod': '∏', '\\coprod': '∐', '\\lim': 'lim',
        '\\limsup': 'lim sup', '\\liminf': 'lim inf', '\\sup': 'sup', '\\inf': 'inf',

        # Delimiters
        '\\langle': '⟨', '\\rangle': '⟩', '\\lceil': '⌈', '\\rceil': '⌉',
        '\\lfloor': '⌊', '\\rfloor': '⌋', '\\lbrack': '[', '\\rbrack': ']',
        '\\lbrace': '{', '\\rbrace': '}', '\\vert': '|', '\\Vert': '‖',
        '\\|': '‖', '\\{': '{', '\\}': '}', '\\(': '(', '\\)': ')',

        # Other
        '\\S': '§', '\\P': '¶', '\\%': '%', '\\$': '$', '\\#': '#', '\\_': '_',
        '\\,': ' ', '\\;': ' ', '\\:': ' ', '\\!': '', '\\quad': ' ', '\\qquad': '  ',
        '\\dots': '…', '\\ldots': '…', '\\cdots': '⋯', '\\vdots': '⋮', '\\ddots': '⋱',
        '\\prime': '′', '\\doubleprime': '″', '\\backprime': '‵', '\\angle': '∠',
        '\\measuredangle': '∡', '\\surd': '√', '\\sqrt': '√', '\\degree': '°',
        '\\aleph': 'ℵ', '\\imath': 'ı', '\\jmath': 'ȷ', '\\ell': 'ℓ', '\\wp': '℘',
        '\\Re': 'ℜ', '\\Im': 'ℑ', '\\mho': '℧', '\\eth': 'ð', '\\gimel': 'ℷ',
        '\\beth': 'ℶ', '\\daleth': 'ℸ', '\\Bbbk': '𝕜',
    }
    
    for latex_symbol, unicode_symbol in symbol_replacements.items():
        latex_expr = latex_expr.replace(latex_symbol, unicode_symbol)
    
    # Clean up remaining LaTeX artifacts
    latex_expr = latex_expr.replace('\\', '')
    latex_expr = latex_expr.replace('{', '').replace('}', '')
    
    # If we haven't converted everything to OMML elements, wrap remaining text
    if not latex_expr.startswith('<m:'):
        latex_expr = f'<m:r><m:t>{latex_expr}</m:t></m:r>'
    
    # Wrap in complete OMML structure
    omml = f'''<m:oMath {nsdecls('m')}>
        {latex_expr}
    </m:oMath>'''
    
    return omml


def insert_math_equation(paragraph, latex_equation):
    """Insert a math equation into a paragraph using OMML conversion"""
    # Create a run for the equation
    run = paragraph.add_run()
    
    # Convert LaTeX to OMML and insert
    try:
        omml_equation = latex_to_omml(latex_equation)
        math_element = parse_xml(omml_equation)
        run._element.append(math_element)
    except Exception as e:
        # Fallback to text if conversion fails
        run.text = latex_equation
        run.font.name = "Cambria Math"
        run.font.italic = True
        print(f"Failed to convert equation to OMML: {str(e)}")


def parse_inline_formatting(text):
    """
    Detect and split text for inline formatting: bold, italic, bold-italic.
    Supports both asterisk and underscore formatting:
    - Bold: **text** or __text__
    - Italic: *text* or _text_
    - Bold-italic: ***text*** or ___text___
    - Bold (alternative): ****text****
    - Inline math: $equation$ or $$equation$$
    """
    # First check for block and inline math expressions
    math_parts = []
    last_math_end = 0
    
    # Match both block math ($$...$$) and inline math ($...$)
    # This pattern handles both single $ and double $$ delimiters
    math_pattern = r'(\$\$(.*?)\$\$|\$((?!\$).*?)(?<!\$)\$)'
    
    for math_match in re.finditer(math_pattern, text):
        math_start, math_end = math_match.span()
        if math_start > last_math_end:
            math_parts.append((text[last_math_end:math_start], {}, False))
        
        equation_text = process_math_equation(math_match.group(0))
        math_parts.append((equation_text, {'math': True}, True))
        last_math_end = math_end
    
    if last_math_end < len(text):
        math_parts.append((text[last_math_end:], {}, False))
    
    # Now process other formatting within non-math parts
    result_parts = []
    for part_text, part_attrs, is_math in math_parts:
        if is_math:
            result_parts.append((part_text, part_attrs))
            continue
            
        pattern = r'(\*\*\*\*.*?\*\*\*\*|\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|\_\_\_.*?\_\_\_|\_\_.*?\_\_|\_.*?\_)'
        last_end = 0
        for match in re.finditer(pattern, part_text):
            start, end = match.span()
            if start > last_end:
                result_parts.append((part_text[last_end:start], {}))
            token = match.group(0)
            format_attrs = {}
            
            # Handle asterisk formats
            if token.startswith('****') and token.endswith('****'):
                content = token[4:-4]
                format_attrs = {'bold': True}
            elif token.startswith('***') and token.endswith('***'):
                content = token[3:-3]
                format_attrs = {'bold': True, 'italic': True}
            elif token.startswith('**') and token.endswith('**'):
                content = token[2:-2]
                format_attrs = {'bold': True}
            elif token.startswith('*') and token.endswith('*'):
                content = token[1:-1]
                format_attrs = {'italic': True}
            # Handle underscore formats
            elif token.startswith('___') and token.endswith('___'):
                content = token[3:-3]
                format_attrs = {'bold': True, 'italic': True}
            elif token.startswith('__') and token.endswith('__'):
                content = token[2:-2]
                format_attrs = {'bold': True}
            elif token.startswith('_') and token.endswith('_'):
                content = token[1:-1]
                format_attrs = {'italic': True}
            else:
                content = token
                
            result_parts.append((content, format_attrs))
            last_end = end
        if last_end < len(part_text):
            result_parts.append((part_text[last_end:], {}))
    
    return result_parts


def add_formatted_text_to_paragraph(paragraph, text, base_font_name='Calibri'):
    """
    Add text with inline formatting to a paragraph.
    """
    for part_text, format_attrs in parse_inline_formatting(text):
        if format_attrs.get('math', False):
            # Handle math expressions with OMML conversion
            insert_math_equation(paragraph, part_text)
        else:
            # Handle regular text formatting
            run = paragraph.add_run(part_text)
            run.font.name = base_font_name
            run.bold = format_attrs.get('bold', False)
            run.italic = format_attrs.get('italic', False)


def add_index_page(doc, toc_lines, primary_color_rgb):
    """
    Add a dedicated Table of Contents page.
    """
    doc.add_page_break()
    toc_heading = doc.add_heading("Table of Contents", level=1)
    toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if toc_heading.runs:
        toc_heading.runs[0].font.name = 'Calibri Light'

    for line in toc_lines:
        match_numbered = re.match(r'^(\d+[\.\)])\s+\[(.*?)\]\(#(.*?)\)', line)
        match_bullet = re.match(r'^(\s*[-*+])\s+\[(.*?)\]\(#(.*?)\)', line)
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)

        link_text_content = ""
        if match_numbered:
            number, link_text_content, _ = match_numbered.groups()
            num_run = p.add_run(f"{number} ") # Corrected space
            num_run.font.name = "Calibri"
            num_run.bold = True
            num_run.font.color.rgb = primary_color_rgb
            p.paragraph_format.left_indent = Inches(0.25)
        elif match_bullet:
            link_text_content, _ = match_bullet.groups()[1:]
            bullet_run = p.add_run("• ")
            bullet_run.font.name = "Calibri"
            bullet_run.font.color.rgb = primary_color_rgb
            p.paragraph_format.left_indent = Inches(0.25)
        else:
            link_text_content = line # Fallback
            p.add_run(line).font.name = "Calibri"
            p.paragraph_format.left_indent = Inches(0.25)
            if line != toc_lines[-1]: doc.add_page_break() # Avoid break after last item
            if not match_numbered and not match_bullet: continue # Skip link creation for fallback

        link_run = p.add_run(link_text_content)
        link_run.font.name = "Calibri"
        link_run.font.color.rgb = primary_color_rgb
        link_run.underline = True

    if toc_lines: # Add page break only if there was a TOC
        doc.add_page_break()


def set_table_style(table, first_row_bg="#005A9C"): # Updated primary color hex
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row in enumerate(table.rows):
        row.height_rule = 1 
        row.height = Pt(24 if i > 0 else 28)

        for cell_idx, cell in enumerate(row.cells):
            set_cell_border(
                cell,
                top={"sz": 1, "val": "single", "color": "#D3D3D3"},
                bottom={"sz": 1, "val": "single", "color": "#D3D3D3"},
                left={"sz": 1, "val": "single", "color": "#D3D3D3"},
                right={"sz": 1, "val": "single", "color": "#D3D3D3"},
            )
            
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            for margin_name, margin_value in [('top', '80'), ('left', '80'), ('bottom', '80'), ('right', '80')]: # Adjusted padding
                node = OxmlElement(f'w:{margin_name}')
                node.set(qn('w:w'), margin_value)
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)
            tcPr.append(tcMar)
            set_cell_vertical_alignment(cell, 'center')
            
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
        
        if i == 0:
            for cell in row.cells:
                set_cell_shading(cell, first_row_bg)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True
                        run.font.name = "Calibri" # Ensure header uses Calibri
                        run.font.size = Pt(10.5)
        # Removed alternating row shading


def process_markdown_table(doc, table_lines):
    """
    Process and create a beautifully formatted table from Markdown table syntax.
    """
    if not table_lines or len(table_lines) < 2: # Need at least header and delimiter
        return
    
    header_line = table_lines[0].strip()
    alignment_line = table_lines[1].strip()
    data_lines = [line.strip() for line in table_lines[2:] if line.strip()]
    
    headers = [cell.strip() for cell in re.split(r'\s*\|\s*', header_line.strip('| '))]   
    align_markers = [m.strip() for m in re.split(r'\s*\|\s*', alignment_line.strip('| '))]   
    alignments = []
    for marker in align_markers:
        if marker.startswith(':') and marker.endswith(':'):
            alignments.append(WD_PARAGRAPH_ALIGNMENT.CENTER)
        elif marker.endswith(':'):
            alignments.append(WD_PARAGRAPH_ALIGNMENT.RIGHT)
        else:
            alignments.append(WD_PARAGRAPH_ALIGNMENT.LEFT)
    
    num_cols = len(headers)
    if num_cols == 0: return
    table = doc.add_table(rows=1 + len(data_lines), cols=num_cols)
    
    # Populate header cells
    for i, header_text in enumerate(headers):
        if i < num_cols:
            cell = table.cell(0, i)
            para = cell.paragraphs[0]
            add_formatted_text_to_paragraph(para, header_text, base_font_name="Calibri")
            if i < len(alignments):
                para.alignment = alignments[i]
    
    # Populate data cells
    for row_idx, row_line in enumerate(data_lines):
        cells_data = [cell.strip() for cell in re.split(r'\s*\|\s*', row_line.strip('| '))]
        for col_idx, cell_text in enumerate(cells_data):
            if col_idx < num_cols:
                cell = table.cell(row_idx + 1, col_idx)
                para = cell.paragraphs[0]
                add_formatted_text_to_paragraph(para, cell_text, base_font_name="Calibri")
                if col_idx < len(alignments):
                    para.alignment = alignments[col_idx]
    
    set_table_style(table)
    doc.add_paragraph("").paragraph_format.space_after = Pt(12) # Space after table


def add_cover_page(doc, title, subtitle=None):
    """
    Add a professional cover page.
    """
    primary_color_hex = "005A9C"
    dark_gray_rgb = RGBColor(0x50, 0x50, 0x50)
    medium_gray_rgb = RGBColor(0x80,0x80,0x80)

    if doc.paragraphs and doc.paragraphs[0].text == "" and not len(doc.paragraphs[0].runs):
        cover_para_for_border = doc.paragraphs[0]
    else:
        cover_para_for_border = doc.add_paragraph()
        if len(doc.element.body):
            doc.element.body.insert(0, cover_para_for_border._p)
        else: # Should not happen if add_paragraph worked
            pass 
    
    cover_para_for_border.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    pPr = cover_para_for_border._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top_border_cover = OxmlElement('w:top') # Changed to top border for visual effect
    top_border_cover.set(qn('w:val'), 'single')
    top_border_cover.set(qn('w:sz'), '12') 
    top_border_cover.set(qn('w:space'), '1')
    top_border_cover.set(qn('w:color'), primary_color_hex)
    pBdr.append(top_border_cover)
    pPr.append(pBdr)
    cover_para_for_border.add_run().add_break() 

    doc.add_paragraph() 
    
    title_heading = doc.add_heading(title, level=0) 
    title_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    if subtitle:
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle_run = subtitle_para.add_run(subtitle)
        subtitle_run.font.name = "Calibri Light"
        subtitle_run.font.size = Pt(18)
        subtitle_run.font.italic = True
        subtitle_run.font.color.rgb = dark_gray_rgb
        subtitle_para.paragraph_format.space_after = Pt(18) # Add space after subtitle
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_run = date_para.add_run(datetime.datetime.now().strftime("%B %d, %Y"))
    date_run.font.name = "Calibri"
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = dark_gray_rgb
    
    for _ in range(8): doc.add_paragraph() # Adjusted spacer
    
    author_para = doc.add_paragraph()
    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author_run = author_para.add_run('Generated by Document Automator')
    author_run.font.name = "Calibri"
    author_run.font.size = Pt(9)
    author_run.font.italic = True
    author_run.font.color.rgb = medium_gray_rgb
    
    doc.add_page_break()


def add_horizontal_rule(doc):
    """
    Add a horizontal rule to the document.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '6')
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), '999999')
    pBdr.append(bottom_border)
    pPr.append(pBdr)


def markdown_to_docx(md_text, output_file, doc_title=None):
    """
    Main conversion function.
    """
    doc = Document()
    set_document_margins(doc)
    code_style_name, blockquote_style_name, math_style_name, primary_color = set_custom_styles(doc)
    
    lines = md_text.splitlines()
    title = doc_title
    subtitle = None
    if not title:
        for line_idx, line_text in enumerate(lines[:5]):
            match = re.match(r'^#\s+(.*)', line_text)
            if match:
                title = match.group(1).strip()
                # Check for subtitle on the next line if it's ##
                if line_idx + 1 < len(lines):
                    subtitle_match = re.match(r'^##\s+(.*)', lines[line_idx+1])
                    if subtitle_match:
                        subtitle = subtitle_match.group(1).strip()
                break
    title = title or "Markdown Document" # Fallback title
    
    # Add cover page first - it might add a paragraph that other logic needs to skip
    add_cover_page(doc, title, subtitle)
    add_decorative_horizontal_line(doc, "D3D3D3", 1, space_before_pt=0, space_after_pt=24) # Line after cover
    
    add_header_footer(doc, title=title, author="Automated Conversion") # Add header/footer after cover page content

    toc_lines = []
    content_lines = []
    in_toc_section = False
    # Skip title/subtitle lines if they were used for cover page from content_lines
    start_content_processing_idx = 0
    if re.match(r'^#\s+(.*)', lines[0] if lines else ""): 
        start_content_processing_idx +=1
        if subtitle and len(lines) > 1 and re.match(r'^##\s+(.*)', lines[1]):
            start_content_processing_idx +=1

    for line_idx in range(start_content_processing_idx, len(lines)):
        line = lines[line_idx]
        if re.match(r'^##\s+Table of Contents', line, re.IGNORECASE):
            in_toc_section = True
            continue
        if in_toc_section:
            if re.match(r'^(\d+[.\)]|[-*+])\s+\[.*?\]\(#.*?\)', line):
                toc_lines.append(line)
                continue
            else:
                in_toc_section = False # End of TOC section
        content_lines.append(line)

    if toc_lines:
        add_index_page(doc, toc_lines, primary_color)
        add_decorative_horizontal_line(doc, "D3D3D3", 1, space_before_pt=0, space_after_pt=18) # Line after TOC

    in_code_block = False
    code_block_lines = []
    in_table_block = False
    table_block_lines = []

    i = 0
    while i < len(content_lines):
        line = content_lines[i]
        
        # Check for standalone block math equation ($$...$$)
        # We only handle standalone equations here; inline equations are handled by the inline formatter
        if re.match(r'^\s*\$\$(.*?)\$\$\s*$', line) and not in_code_block:
            equation_text = process_math_equation(line)
            p = doc.add_paragraph(style=math_style_name)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # Use OMML to render the equation
            insert_math_equation(p, equation_text)
            i += 1
            continue
            
        # Check for horizontal rule
        if re.match(r'^(\s*)(---|\*\*\*|___)(\s*)$', line):
            add_horizontal_rule(doc)
            i += 1
            continue
            
        if re.match(r'^(\s*```)', line): # Code block toggle
            in_code_block = not in_code_block
            if not in_code_block and code_block_lines: # End of code block
                p = doc.add_paragraph("\n".join(code_block_lines), style=code_style_name)
                apply_code_block_formatting(p)
                code_block_lines = []
            i += 1
            continue
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Table processing
        is_table_line = "|" in line
        is_table_delimiter = re.match(r'^(\s*)\|[\s\-:]+\|', line)
        
        if not in_table_block and is_table_line and i + 1 < len(content_lines) and re.match(r'^(\s*)\|[\s\-:]+\|', content_lines[i+1]):
            in_table_block = True
            table_block_lines.append(line) # Header row
            i += 1
            continue
        
        if in_table_block:
            if is_table_line:
                table_block_lines.append(line)
                i += 1
                # Check if this is the last line or next line is not part of table
                if i == len(content_lines) or not ("|" in content_lines[i]): 
                    process_markdown_table(doc, table_block_lines)
                    table_block_lines = []
                    in_table_block = False
                    # Current line (content_lines[i]) is not part of table, will be processed in next iteration
                continue 
            else: # Line is not part of table, so table ended
                process_markdown_table(doc, table_block_lines)
                table_block_lines = []
                in_table_block = False
                # Current line (line) will be processed normally below
        
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            heading = doc.add_heading('', level=level)
            add_formatted_text_to_paragraph(heading, text, base_font_name="Calibri Light")
            if level == 1 and i > 0 : # Add line after H1, but not if it's the very first content item
                 # Check if previous line was also H1, to avoid double lines if H1s are consecutive
                is_first_content_h1 = True
                if i-1 >= 0 and re.match(r'^(#{1,6})\s+(.*)', content_lines[i-1]):
                    pass # Potentially avoid if prev was also a heading, or fine-tune logic
                else: # Only add if not immediately preceded by another heading, or always add for H1
                    add_decorative_horizontal_line(doc, "E0E0E0", 0.75, space_before_pt=4, space_after_pt=10)
            i += 1
            continue

        blockquote_match = re.match(r'^(\s*>+)\s?(.*)', line)
        if blockquote_match:
            text = blockquote_match.group(2).strip()
            p = doc.add_paragraph(style=blockquote_style_name) 
            add_formatted_text_to_paragraph(p, text, base_font_name="Calibri")
            i += 1
            continue

        # Ordered list (numbered)
        ordered_list_match = re.match(r'^(\s*)(\d+[.)])\s+(.*)', line)
        if ordered_list_match:
            leading_spaces_str = ordered_list_match.group(1)
            list_marker = ordered_list_match.group(2)
            text = ordered_list_match.group(3).strip()
            
            # Calculate nesting level (normalize tabs to spaces first)
            leading_spaces = len(leading_spaces_str.expandtabs(INDENT_SPACES_PER_LEVEL))
            level = leading_spaces // INDENT_SPACES_PER_LEVEL
            
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = BASE_LIST_INDENT + (LIST_INDENT_INCREMENT * level)
            p.paragraph_format.first_line_indent = HANGING_LIST_INDENT
            p.paragraph_format.space_after = Pt(4)
            
            add_formatted_text_to_paragraph(p, text, base_font_name="Calibri")
            i += 1
            continue

        # Unordered list (bullet)
        list_match = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if list_match:
            leading_spaces_str = list_match.group(1)
            text = list_match.group(2).strip()
            
            # Calculate nesting level (normalize tabs to spaces first)
            leading_spaces = len(leading_spaces_str.expandtabs(INDENT_SPACES_PER_LEVEL))
            level = leading_spaces // INDENT_SPACES_PER_LEVEL
            
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = BASE_LIST_INDENT + (LIST_INDENT_INCREMENT * level)
            p.paragraph_format.first_line_indent = HANGING_LIST_INDENT
            p.paragraph_format.space_after = Pt(4) 
            
            add_formatted_text_to_paragraph(p, text, base_font_name="Calibri")
            i += 1
            continue

        if not line.strip(): # Blank line
            doc.add_paragraph("")
            i += 1
            continue

        # Default: regular paragraph
        p = doc.add_paragraph()
        add_formatted_text_to_paragraph(p, line.strip(), base_font_name="Calibri")
        i += 1

    # Process any remaining code block or table at the end of the document
    if in_code_block and code_block_lines:
        p = doc.add_paragraph("\n".join(code_block_lines), style=code_style_name)
        apply_code_block_formatting(p)
    if in_table_block and table_block_lines:
        process_markdown_table(doc, table_block_lines)

    core_props = doc.core_properties
    core_props.title = title
    core_props.author = "Document Automator Script"
    core_props.comments = f"Generated from Markdown on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"

    doc.save(output_file)
    print(f"Document saved as {output_file}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert Markdown to DOCX with enhanced styling.")
    parser.add_argument("input_file", help="Path to the input Markdown file")
    parser.add_argument("output_file", nargs="?", help="Path to the output DOCX file (optional)")
    parser.add_argument("--title", help="Document title (optional, overrides title from Markdown)")

    args = parser.parse_args()

    output_file_path = args.output_file
    if not output_file_path:
        output_file_path = os.path.splitext(args.input_file)[0] + ".docx"

    if not os.path.exists(args.input_file):
        print(f"Error: Input file '{args.input_file}' does not exist.", file=sys.stderr)
        sys.exit(1)

    try:
        with open(args.input_file, 'r', encoding='utf-8') as f:
            md_text_content = f.read()
    except Exception as e:
        print(f"Error reading input file '{args.input_file}': {e}", file=sys.stderr)
        sys.exit(1)

    try:
        markdown_to_docx(md_text_content, output_file_path, args.title)
    except Exception as e:
        print(f"Error during Markdown to DOCX conversion: {e}", file=sys.stderr)
        # Consider adding more detailed error logging or traceback for debugging
        import traceback
        print(traceback.format_exc(), file=sys.stderr)
        sys.exit(1)